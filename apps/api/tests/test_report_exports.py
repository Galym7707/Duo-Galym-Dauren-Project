from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.api import routes
from app.main import app
from app.providers.gee import GeeCandidate, GeeSyncSummary
from app.services.workflow_store import WorkflowStore


def make_client() -> TestClient:
    routes.replace_runtime_services(WorkflowStore())
    return TestClient(app)


def make_live_candidate() -> GeeCandidate:
    return GeeCandidate(
        id="GEE-20260327-01",
        asset_name="Atyrau Region CH4 hotspot 01",
        region="Atyrau Region",
        facility_type="Methane hotspot with night thermal context",
        severity="high",
        detected_at="2026-03-27 08:00",
        methane_delta_pct=3.41,
        methane_delta_ppb=62.2,
        signal_score=82,
        confidence="High screening confidence / methane uplift plus night thermal context",
        coordinates="46.190 N, 51.858 E",
        latitude=46.19,
        longitude=51.858,
        summary="Live candidate summary",
        recommended_action="Promote this candidate into an incident and send it to field verification.",
        current_ch4_ppb=1884.6,
        baseline_ch4_ppb=1822.4,
        thermal_hits_72h=12,
        night_thermal_hits_72h=12,
        evidence_source="Google Earth Engine / Sentinel-5P + VIIRS thermal context",
        baseline_window="84-day Kazakhstan baseline before 2026-03-27 08:00 UTC",
        verification_area="Makat District, Atyrau Region",
        nearest_address="A27, Atyrau Region",
        nearest_landmark="Tengiz Field",
    )


def prepare_export(client: TestClient, locale: str = "en"):
    routes.pipeline_service.provider.sync_summary = lambda: GeeSyncSummary(
        project_id="demo-project",
        status="ready",
        message="Earth Engine CH4 screening summary fetched successfully.",
        latest_observation_at="2026-03-27 08:00 UTC",
        observed_window="Latest TROPOMI scene compared with Kazakhstan historical mean.",
        mean_ch4_ppb=1884.6,
        baseline_ch4_ppb=1822.4,
        delta_abs_ppb=62.2,
        delta_pct=3.41,
        scene_count=12,
        candidates=[make_live_candidate()],
    )
    sync_response = client.post("/api/v1/pipeline/sync", json={"source": "gee"})
    assert sync_response.status_code == 200
    dashboard = client.get("/api/v1/dashboard")
    anomaly_id = dashboard.json()["anomalies"][0]["id"]
    promote = client.post(f"/api/v1/anomalies/{anomaly_id}/promote", json={"owner": "ESG desk"})
    assert promote.status_code == 201
    incident_id = promote.json()["id"]
    report = client.post(f"/api/v1/incidents/{incident_id}/report")
    assert report.status_code == 200
    return routes.store._prepare_report_export(incident_id, locale)


def test_prepare_report_exposes_management_and_timeline_blocks() -> None:
    client = make_client()
    report = prepare_export(client, "en")

    assert report.management_title == "Management snapshot"
    assert report.evidence_title == "Operational evidence"
    assert report.recommended_action_title == "Recommended action"
    assert len(report.hero_metrics) == 3
    assert len(report.supporting_metrics) == 3
    assert any(label == "Document ID" and value.startswith("INC-") for label, value in report.metadata)
    assert report.audit_timeline
    assert report.audit_timeline[0].stage_label == "Screening"


def test_renderers_share_executive_first_structure() -> None:
    client = make_client()
    report = prepare_export(client, "en")
    incident_id = next(value for label, value in report.metadata if label == "Document ID")

    html = routes.store.export_report_html(incident_id, locale="en")
    pdf = routes.store.export_report_pdf(incident_id, locale="en")
    docx = routes.store.export_report_docx(incident_id, locale="en")

    assert "Management snapshot" in html
    assert "Recommended action" in html
    assert "Operational evidence" in html
    assert "Document context" in html
    assert "MRV report for screening and operational prioritization" in html
    assert pdf.startswith(b"%PDF")

    document = Document(BytesIO(docx))
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Management snapshot" in full_text
    assert "Recommended action" in full_text
    assert "Operational evidence" in full_text


def test_russian_export_keeps_new_structure_labels_readable() -> None:
    client = make_client()
    report = prepare_export(client, "ru")
    incident_id = next(value for label, value in report.metadata if label == "ID документа")

    html = routes.store.export_report_html(incident_id, locale="ru")
    docx = routes.store.export_report_docx(incident_id, locale="ru")

    assert "Управленческий срез" in html
    assert "Рекомендуемое действие" in html
    assert "Операционные доказательства" in html

    document = Document(BytesIO(docx))
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Управленческий срез" in full_text
    assert "Рекомендуемое действие" in full_text
