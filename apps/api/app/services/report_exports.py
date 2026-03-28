from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models import ActivityEvent, Anomaly, Incident, ReportSection

Locale = Literal["en", "ru"]


@dataclass(frozen=True)
class ReportTaskLine:
    title: str
    owner: str
    eta: str
    status: str
    notes: str


@dataclass(frozen=True)
class PreparedReport:
    locale: Locale
    file_stem: str
    title: str
    subtitle: str
    metadata: list[tuple[str, str]]
    sections: list[ReportSection]
    tasks_title: str
    tasks: list[ReportTaskLine]
    audit_title: str
    audit_lines: list[str]


REGION_TRANSLATIONS = {
    "Atyrau Region": "Атырауская область",
    "Mangystau Region": "Мангистауская область",
    "Aktobe Region": "Актюбинская область",
    "West Kazakhstan Region": "Западно-Казахстанская область",
    "Kyzylorda Region": "Кызылординская область",
    "Pavlodar Region": "Павлодарская область",
    "Akmola Region": "Акмолинская область",
    "Almaty Region": "Алматинская область",
    "Almaty City": "Алматы",
    "Karaganda Region": "Карагандинская область",
    "Kostanay Region": "Костанайская область",
    "North Kazakhstan Region": "Северо-Казахстанская область",
    "East Kazakhstan Region": "Восточно-Казахстанская область",
    "Turkistan Region": "Туркестанская область",
    "Zhambyl Region": "Жамбылская область",
    "Kazakhstan": "Казахстан",
}

ASSET_TRANSLATIONS = {
    "Tengiz satellite cluster": "Тенгизский спутниковый кластер",
    "Karabatan processing block": "Карабатанский перерабатывающий блок",
    "Mangystau export hub": "Мангистауский экспортный узел",
    "Aktobe compressor ring": "Актюбинское компрессорное кольцо",
    "Karachaganak gas train": "Карачаганакская газовая нитка",
    "Kumkol gathering node": "Кумкольский узел сбора",
    "Pavlodar refinery corridor": "Павлодарский перерабатывающий коридор",
}

OWNER_TRANSLATIONS = {
    "Field integrity desk": "Группа полевой целостности",
    "Ops coordinator": "Координатор эксплуатации",
    "Reliability engineer": "Инженер по надежности",
    "ESG lead": "Руководитель ESG",
    "Response lead": "Ответственный за реагирование",
    "Remote sensing analyst": "Аналитик дистанционного зондирования",
    "Area operations coordinator": "Координатор площадки",
    "Compliance lead": "Руководитель по соответствию",
    "ESG desk": "ESG-команда",
    "MRV response lead": "Координатор MRV-реагирования",
}

TASK_TITLE_TRANSLATIONS = {
    "Dispatch LDAR walkdown request": "Отправить запрос на LDAR-обход",
    "Cross-check flare line maintenance history": "Проверить историю обслуживания факельной линии",
    "Draft regulator-facing MRV note": "Подготовить MRV-заметку для регулятора",
    "Validate signal persistence against 12-week baseline": "Проверить устойчивость сигнала по периоду в 12 недель",
    "Assign field verification owner": "Назначить ответственного за выездную проверку",
    "Collect operator comment": "Собрать комментарий оператора",
}

AUDIT_SOURCE_TRANSLATIONS = {
    "seeded": "Демо-источник",
    "gee": "Google Earth Engine",
    "workflow": "Рабочий процесс",
}


def prepare_report(
    anomaly: Anomaly,
    incident: Incident,
    audit_events: list[ActivityEvent],
    locale: Locale = "en",
) -> PreparedReport:
    completed_tasks = sum(1 for task in incident.tasks if task.status == "done")
    labels = _labels(locale)

    metadata = [
        (labels["generated"], incident.report_generated_at or labels["on_demand"]),
        (labels["incident"], incident.id),
        (labels["asset"], _translate_asset(anomaly.asset_name, locale)),
        (labels["region"], _translate_region(anomaly.region, locale)),
        (labels["coordinates"], anomaly.coordinates),
    ]
    if anomaly.verification_area:
        metadata.append((labels["verification_area"], anomaly.verification_area))
    if anomaly.nearest_address:
        metadata.append((labels["nearest_address"], anomaly.nearest_address))
    if anomaly.nearest_landmark:
        metadata.append((labels["nearest_landmark"], anomaly.nearest_landmark))
    metadata.extend(
        [
            (labels["owner"], _translate_owner(incident.owner, locale)),
            (labels["priority"], incident.priority),
            (labels["window"], _translate_window(incident.verification_window, locale)),
            (_impact_label(anomaly, locale), _impact_value(anomaly, locale)),
            (labels["progress"], _task_progress(completed_tasks, len(incident.tasks), locale)),
        ]
    )

    tasks = [
        ReportTaskLine(
            title=_translate_task_title(task.title, locale),
            owner=_translate_owner(task.owner, locale),
            eta=_format_hours(task.eta_hours, locale),
            status=labels["done"] if task.status == "done" else labels["open"],
            notes=task.notes,
        )
        for task in incident.tasks
    ]

    return PreparedReport(
        locale=locale,
        file_stem=f"{incident.id.lower()}-mrv-report",
        title=f"{labels['title']}: {incident.id}",
        subtitle=labels["subtitle"],
        metadata=metadata,
        sections=_localized_sections(anomaly, incident, completed_tasks, locale),
        tasks_title=labels["tasks"],
        tasks=tasks,
        audit_title=labels["audit"],
        audit_lines=[_format_audit_line(event, locale) for event in audit_events],
    )


def render_html(report: PreparedReport, auto_print: bool = False) -> str:
    labels = _labels(report.locale)
    meta_markup = "".join(
        f"<div><span class='label'>{_escape_html(label)}</span><span class='value'>{_escape_html(value)}</span></div>"
        for label, value in report.metadata
    )
    section_markup = "".join(
        f"<section><h2>{_escape_html(section.title)}</h2><p>{_escape_html(section.body)}</p></section>"
        for section in report.sections
    )
    task_markup = "".join(
        (
            "<li>"
            f"<strong>{_escape_html(task.title)}</strong> - {_escape_html(task.owner)} - {_escape_html(task.eta)}"
            f" - {_escape_html(task.status)}"
            f"{f'<br /><small>{_escape_html(task.notes)}</small>' if task.notes else ''}"
            "</li>"
        )
        for task in report.tasks
    )
    audit_markup = "".join(f"<li>{_escape_html(line)}</li>" for line in report.audit_lines)
    auto_print_script = (
        "<script>window.addEventListener('load', () => { setTimeout(() => window.print(), 120); });</script>"
        if auto_print
        else ""
    )

    return (
        "<!doctype html>"
        f"<html lang='{report.locale}'>"
        "<head>"
        "<meta charset='utf-8' />"
        "<meta name='viewport' content='width=device-width, initial-scale=1' />"
        f"<title>{_escape_html(report.title)}</title>"
        "<style>"
        "body{font-family:Segoe UI,Arial,sans-serif;margin:40px;color:#10212b;line-height:1.55;background:#f5efe6;}"
        "main{max-width:920px;margin:0 auto;padding:36px;border:1px solid #d7d0c4;background:#fffdfa;}"
        "h1{margin:0 0 8px;font-size:32px;}h2{margin:24px 0 8px;font-size:20px;}section{margin-top:20px;}"
        ".meta{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px 20px;margin:24px 0;}"
        ".meta div{padding:12px 14px;border:1px solid #d3dde5;background:#f6fafc;}"
        ".label{display:block;font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#5c6f7b;}"
        ".value{display:block;margin-top:6px;font-weight:600;}"
        ".toolbar{display:flex;gap:12px;align-items:center;justify-content:space-between;margin:20px 0 28px;}"
        ".toolbar button{padding:10px 14px;border:1px solid #b9c9d4;background:#ffffff;cursor:pointer;font:inherit;}"
        "ul{padding-left:20px;}li{margin-bottom:8px;}small{color:#5c6f7b;}"
        "@media print{body{margin:22px;background:#fff;}main{padding:0;border:none;background:#fff;} .toolbar{display:none;} .meta{gap:8px 14px;}}"
        "</style>"
        "</head>"
        "<body>"
        "<main>"
        f"<h1>{_escape_html(report.title)}</h1>"
        f"<p>{_escape_html(report.subtitle)}</p>"
        "<div class='toolbar'>"
        f"<span>{_escape_html(report.subtitle)}</span>"
        f"<button onclick='window.print()'>{_escape_html(labels['print'])}</button>"
        "</div>"
        f"<div class='meta'>{meta_markup}</div>"
        f"{section_markup}"
        f"<section><h2>{_escape_html(report.tasks_title)}</h2><ul>{task_markup}</ul></section>"
        f"<section><h2>{_escape_html(report.audit_title)}</h2><ul>{audit_markup}</ul></section>"
        f"{auto_print_script}"
        "</main>"
        "</body></html>"
    )


def render_pdf(report: PreparedReport) -> bytes:
    regular_font, bold_font = _resolve_pdf_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontName=bold_font,
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#10212b"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#4b5762"),
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "ReportSection",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#10212b"),
        spaceBefore=14,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#33434f"),
    )

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    elements = [
        Paragraph(_escape_html(report.title), title_style),
        Paragraph(_escape_html(report.subtitle), subtitle_style),
    ]

    meta_rows = []
    for index in range(0, len(report.metadata), 2):
        cells = report.metadata[index:index + 2]
        if len(cells) == 1:
            cells = [cells[0], ("", "")]
        meta_rows.append(
            [
                Paragraph(f"<b>{_escape_html(label)}</b><br/>{_escape_html(value)}", body_style)
                for label, value in cells
            ]
        )
    meta_table = Table(meta_rows, colWidths=[83 * mm, 83 * mm], hAlign="LEFT")
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f6fafc")),
                ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#d3dde5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d3dde5")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    elements.extend([meta_table, Spacer(1, 10)])

    for section in report.sections:
        elements.append(Paragraph(_escape_html(section.title), section_style))
        elements.append(Paragraph(_escape_html(section.body), body_style))

    elements.append(Paragraph(_escape_html(report.tasks_title), section_style))
    for task in report.tasks:
        task_line = (
            f"<b>{_escape_html(task.title)}</b> - {_escape_html(task.owner)} - {_escape_html(task.eta)}"
            f" - {_escape_html(task.status)}"
        )
        elements.append(Paragraph(task_line, body_style))
        if task.notes:
            elements.append(Paragraph(_escape_html(task.notes), body_style))
        elements.append(Spacer(1, 4))

    elements.append(Paragraph(_escape_html(report.audit_title), section_style))
    for line in report.audit_lines:
        elements.append(Paragraph(_escape_html(line), body_style))
        elements.append(Spacer(1, 4))

    document.build(elements)
    return buffer.getvalue()


def render_docx(report: PreparedReport) -> bytes:
    document = Document()
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11 if style_name == "Normal" else 18 if style_name == "Heading 1" else 13)

    title = document.add_heading(report.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(report.subtitle)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in report.metadata:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    for section in report.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.body)

    document.add_heading(report.tasks_title, level=2)
    for task in report.tasks:
        item = document.add_paragraph(style="List Bullet")
        item.add_run(task.title).bold = True
        item.add_run(f" - {task.owner} - {task.eta} - {task.status}")
        if task.notes:
            note = document.add_paragraph(task.notes)
            note.paragraph_format.left_indent = Pt(18)

    document.add_heading(report.audit_title, level=2)
    for line in report.audit_lines:
        document.add_paragraph(line, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _labels(locale: Locale) -> dict[str, str]:
    if locale == "ru":
        return {
            "title": "Отчет Saryna MRV",
            "subtitle": "Краткий MRV-отчет по текущему кейсу со спутниковым скринингом метана.",
            "generated": "Сформирован",
            "on_demand": "По запросу",
            "incident": "Инцидент",
            "asset": "Зона",
            "region": "Регион",
            "coordinates": "Координаты",
            "verification_area": "Район проверки",
            "nearest_address": "Ближайший адрес",
            "nearest_landmark": "Ближайший ориентир",
            "owner": "Ответственный",
            "priority": "Приоритет",
            "window": "Срок проверки",
            "impact": "Потенциальный эффект",
            "progress": "Прогресс",
            "tasks": "Задачи проверки",
            "audit": "История действий",
            "done": "Выполнено",
            "open": "Открыто",
            "print": "Печать / Сохранить как PDF",
        }

    return {
        "title": "Saryna MRV Report",
        "subtitle": "Short MRV note for the current methane screening case.",
        "generated": "Generated",
        "on_demand": "On-demand",
        "incident": "Incident",
        "asset": "Area",
        "region": "Region",
        "coordinates": "Coordinates",
        "verification_area": "Verification area",
        "nearest_address": "Nearest address",
        "nearest_landmark": "Nearest landmark",
        "owner": "Owner",
        "priority": "Priority",
        "window": "Verification window",
        "impact": "Potential impact",
        "progress": "Progress",
        "tasks": "Verification tasks",
        "audit": "Audit Timeline",
        "done": "Done",
        "open": "Open",
        "print": "Print / Save as PDF",
    }


def _localized_sections(
    anomaly: Anomaly,
    incident: Incident,
    completed_tasks: int,
    locale: Locale,
) -> list[ReportSection]:
    return [
        ReportSection(
            title="Что увидели" if locale == "ru" else "What was observed",
            body=_measurement_section_text(anomaly, locale),
        ),
        ReportSection(
            title="Кто отвечает" if locale == "ru" else "Who owns the case",
            body=_ownership_section_text(incident, locale),
        ),
        ReportSection(
            title="Как продвигается проверка" if locale == "ru" else "How verification is progressing",
            body=_progress_section_text(anomaly, incident, completed_tasks, locale),
        ),
    ]


def _impact_label(anomaly: Anomaly, locale: Locale) -> str:
    if anomaly.co2e_tonnes is not None:
        return _labels(locale)["impact"]
    return "Термоконтекст (72 часа)" if locale == "ru" else "Thermal context (72h)"


def _impact_value(anomaly: Anomaly, locale: Locale) -> str:
    if anomaly.co2e_tonnes is not None:
        return f"{anomaly.co2e_tonnes} tCO2e"
    hits = anomaly.night_thermal_hits_72h or anomaly.thermal_hits_72h or 0
    if locale == "ru":
        return f"{hits} ночных VIIRS-срабатываний"
    return f"{hits} night-time VIIRS detections"


def _measurement_section_text(anomaly: Anomaly, locale: Locale) -> str:
    translated_asset = _translate_asset(anomaly.asset_name, locale)
    translated_region = _translate_region(anomaly.region, locale)
    delta_pct = f"{anomaly.methane_delta_pct:.2f}"

    if anomaly.co2e_tonnes is not None:
        flare_hours = anomaly.flare_hours or 0
        if locale == "ru":
            return (
                f"Спутниковый разбор отметил объект {translated_asset} в регионе {translated_region} "
                f"с ростом метана на {delta_pct}% и {flare_hours} часами факельной активности."
            )
        return (
            f"Satellite review flagged {translated_asset} in {translated_region} with {delta_pct}% "
            f"methane uplift and {flare_hours} hours of flare activity."
        )

    delta_ppb = anomaly.methane_delta_ppb or 0
    night_hits = anomaly.night_thermal_hits_72h or 0
    if locale == "ru":
        if night_hits > 0:
            return (
                f"Живой спутниковый скрининг отметил зону {translated_asset} в регионе {translated_region} "
                f"с ростом метана на {delta_ppb:.2f} ppb ({delta_pct}%) и {night_hits} ночными "
                "VIIRS-срабатываниями в радиусе 25 км."
            )
        return (
            f"Живой спутниковый скрининг отметил зону {translated_asset} в регионе {translated_region} "
            f"с ростом метана на {delta_ppb:.2f} ppb ({delta_pct}%) без недавних ночных "
            "VIIRS-срабатываний в радиусе 25 км."
        )
    if night_hits > 0:
        return (
            f"Live satellite screening flagged {translated_asset} in {translated_region} with "
            f"{delta_ppb:.2f} ppb ({delta_pct}%) methane uplift and {night_hits} night-time VIIRS detections "
            "inside the 25 km context window."
        )
    return (
        f"Live satellite screening flagged {translated_asset} in {translated_region} with "
        f"{delta_ppb:.2f} ppb ({delta_pct}%) methane uplift and no recent night-time VIIRS detections "
        "inside the 25 km context window."
    )


def _ownership_section_text(incident: Incident, locale: Locale) -> str:
    translated_owner = _translate_owner(incident.owner, locale)
    translated_window = _translate_window(incident.verification_window, locale).lower()
    if locale == "ru":
        return (
            f"{translated_owner} ведет этот кейс с приоритетом {incident.priority} "
            f"и сроком реакции {translated_window}."
        )
    return (
        f"{translated_owner} owns this case under {incident.priority} priority "
        f"with a {translated_window} response window."
    )


def _progress_section_text(
    anomaly: Anomaly,
    incident: Incident,
    completed_tasks: int,
    locale: Locale,
) -> str:
    if anomaly.co2e_tonnes is not None:
        if locale == "ru":
            return (
                f"Выполнено {completed_tasks} из {len(incident.tasks)} задач. "
                f"Текущая оценка возможного эффекта составляет {anomaly.co2e_tonnes} tCO2e."
            )
        return (
            f"{completed_tasks} of {len(incident.tasks)} tasks are complete. "
            f"The current estimated impact is {anomaly.co2e_tonnes} tCO2e."
        )
    if locale == "ru":
        return (
            f"Выполнено {completed_tasks} из {len(incident.tasks)} задач. "
            "Живая очередь пока не переводит этот сигнал в tCO2e: она честно ранжирует кейсы "
            "по росту метана и тепловому контексту рядом с точкой наблюдения."
        )
    return (
        f"{completed_tasks} of {len(incident.tasks)} tasks are complete. "
        "This live queue does not estimate tCO2e yet; it ranks the case by methane uplift and nearby thermal context."
    )


def _format_audit_line(event: ActivityEvent, locale: Locale) -> str:
    source = AUDIT_SOURCE_TRANSLATIONS.get(event.source, event.source) if locale == "ru" else event.source
    if locale == "ru":
        return f"{event.occurred_at}: {event.title}. {event.detail} Источник: {source}. Исполнитель: {event.actor}."
    return f"{event.occurred_at}: {event.title}. {event.detail} Source: {source}. Actor: {event.actor}."


def _translate_region(value: str, locale: Locale) -> str:
    return REGION_TRANSLATIONS.get(value, value) if locale == "ru" else value


def _translate_asset(value: str, locale: Locale) -> str:
    return ASSET_TRANSLATIONS.get(value, value) if locale == "ru" else value


def _translate_owner(value: str, locale: Locale) -> str:
    return OWNER_TRANSLATIONS.get(value, value) if locale == "ru" else value


def _translate_task_title(value: str, locale: Locale) -> str:
    return TASK_TITLE_TRANSLATIONS.get(value, value) if locale == "ru" else value


def _translate_window(value: str, locale: Locale) -> str:
    if locale == "ru":
        if value == "Next 12 hours":
            return "в ближайшие 12 часов"
        if value == "Next 24 hours":
            return "в ближайшие 24 часа"
        if value == "Next 48 hours":
            return "в ближайшие 48 часов"
    return value


def _task_progress(done: int, total: int, locale: Locale) -> str:
    return f"{done} из {total}" if locale == "ru" else f"{done} of {total}"


def _format_hours(hours: int, locale: Locale) -> str:
    return f"{hours} ч" if locale == "ru" else f"{hours}h"


def _escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def _resolve_pdf_fonts() -> tuple[str, str]:
    font_candidates = [
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            "ArialUnicode",
            "ArialUnicodeBold",
        ),
        (
            Path("C:/Windows/Fonts/DejaVuSans.ttf"),
            Path("C:/Windows/Fonts/DejaVuSans-Bold.ttf"),
            "DejaVuSans",
            "DejaVuSansBold",
        ),
    ]

    for regular_path, bold_path, regular_name, bold_name in font_candidates:
        if regular_path.exists() and bold_path.exists():
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
            return regular_name, bold_name

    return "Helvetica", "Helvetica-Bold"
